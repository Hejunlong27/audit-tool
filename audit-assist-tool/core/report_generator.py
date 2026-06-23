# -*- coding: utf-8 -*-
"""
审计报告生成器
基于 Jinja2 模板 + python-docx 生成审计报告
"""
import os
from typing import Optional
from loguru import logger

from core.base import BaseProcessor, ProcessResult


class ReportGenerator(BaseProcessor):
    """审计报告生成器"""

    def process(
        self,
        template_path: str,
        output_path: str,
        data: dict
    ) -> ProcessResult:
        """
        基于模板生成审计报告

        Args:
            template_path: Word 模板路径 (.docx)
            output_path: 输出文件路径
            data: 填充数据字典
        """
        result = ProcessResult()

        try:
            from docxtpl import DocxTemplate

            if not os.path.exists(template_path):
                result.add_error(f"模板文件不存在: {template_path}")
                return result

            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

            doc = DocxTemplate(template_path)
            doc.render(data)
            doc.save(output_path)

            file_size = os.path.getsize(output_path)
            result.data = {'output_path': output_path, 'file_size': file_size}
            result.message = f"报告生成完成: {output_path} ({file_size/1024:.1f} KB)"
            logger.info(result.message)

        except ImportError:
            result.add_error("docxtpl 未安装，请运行: pip install docxtpl")
        except Exception as e:
            result.add_error(f"报告生成失败: {str(e)}")
            logger.exception(e)

        return result

    def generate_from_data(
        self,
        output_path: str,
        title: str,
        client_name: str,
        audit_period: str,
        findings: list[dict],
        conclusion: str
    ) -> ProcessResult:
        """
        快速生成审计报告（使用内置模板）

        Args:
            output_path: 输出路径
            title: 报告标题
            client_name: 被审计单位
            audit_period: 审计期间
            findings: 审计发现列表 [{"issue": "", "risk": "", "suggestion": ""}]
            conclusion: 审计结论
        """
        data = {
            'title': title,
            'client_name': client_name,
            'audit_period': audit_period,
            'findings': findings,
            'conclusion': conclusion,
            'report_date': '2026年6月23日'
        }

        # 查找内置模板
        template_dir = self.config.get('template_dir', './templates/audit_reports')
        template_path = os.path.join(template_dir, 'standard_report.docx')

        if not os.path.exists(template_path):
            # 如果没有模板，使用 python-docx 直接创建
            return self._create_simple_report(output_path, data)

        return self.process(template_path, output_path, data)

    def _create_simple_report(self, output_path: str, data: dict) -> ProcessResult:
        """无模板时直接创建简单报告"""
        result = ProcessResult()

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 标题
            title = doc.add_heading(data['title'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本信息
            doc.add_paragraph(f"被审计单位: {data['client_name']}")
            doc.add_paragraph(f"审计期间: {data['audit_period']}")
            doc.add_paragraph(f"报告日期: {data['report_date']}")

            doc.add_heading('审计发现', level=1)
            for i, finding in enumerate(data.get('findings', []), 1):
                doc.add_heading(f"发现 {i}: {finding.get('issue', '')}", level=2)
                doc.add_paragraph(f"风险等级: {finding.get('risk', '')}")
                doc.add_paragraph(f"建议: {finding.get('suggestion', '')}")

            doc.add_heading('审计结论', level=1)
            doc.add_paragraph(data.get('conclusion', ''))

            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
            doc.save(output_path)

            result.data = {'output_path': output_path}
            result.message = f"报告生成完成: {output_path}"

        except Exception as e:
            result.add_error(f"报告创建失败: {str(e)}")

        return result
